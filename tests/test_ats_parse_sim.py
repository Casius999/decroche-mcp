"""Tests for ats.parse_sim — structural ATS simulation.

TDD: written before implementation.
"""
from __future__ import annotations

from pathlib import Path

import pytest


# ── Fixtures ───────────────────────────────────────────────────────────────

def make_single_col_pdf(path: Path) -> None:
    """Single-column PDF with selectable text (all text in left region x < 300)."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=LETTER)
    y = 750
    lines = [
        "Jane Doe",
        "jane.doe@example.com",
        "",
        "Experience",
        "Led migration of 12 services to Kubernetes",
        "Reduced API latency 38% by introducing a caching layer",
        "",
        "Education",
        "B.S. Computer Science, MIT",
        "",
        "Skills",
        "Python, Go, Kubernetes, PostgreSQL",
    ]
    for line in lines:
        c.drawString(50, y, line)
        y -= 14
    c.save()


def make_two_col_pdf(path: Path) -> None:
    """Two-column PDF: text placed in two distinct x-bands (left and right)."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=LETTER)
    # Left column
    for i, text in enumerate(["Name", "Email", "Phone", "Address", "LinkedIn"]):
        c.drawString(50, 750 - i * 14, text)
    # Right column — far right to ensure two clusters
    for i, text in enumerate(["Experience", "Led team", "Education", "MIT", "Skills"]):
        c.drawString(350, 750 - i * 14, text)
    c.save()


def make_table_docx(path: Path) -> None:
    """DOCX with a table (skills in a table)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Led migration of 12 services")

    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Go"
    table.cell(0, 2).text = "Kubernetes"
    table.cell(1, 0).text = "PostgreSQL"
    table.cell(1, 1).text = "Redis"
    table.cell(1, 2).text = "AWS"
    doc.save(str(path))


def make_header_contact_docx(path: Path) -> None:
    """DOCX with contact info (email/phone) in the header."""
    from docx import Document

    doc = Document()
    # Add contact to header
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "jane.doe@example.com | +1 415 555 0101"

    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Led migration of 12 services to Kubernetes")
    doc.add_paragraph("Education")
    doc.add_paragraph("MIT Computer Science")
    doc.save(str(path))


@pytest.fixture(scope="session")
def ats_fixtures_dir(tmp_path_factory) -> Path:
    d = tmp_path_factory.mktemp("ats_fixtures")
    make_single_col_pdf(d / "single_col.pdf")
    make_two_col_pdf(d / "two_col.pdf")
    make_table_docx(d / "table.docx")
    make_header_contact_docx(d / "header_contact.docx")
    return d


# ── Tests ─────────────────────────────────────────────────────────────────

def test_single_col_high_parsability_workday(ats_fixtures_dir: Path) -> None:
    """Single-column, clean PDF → high parsability for workday."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "single_col.pdf", "workday")
    assert result.ats_id == "workday"
    assert result.parsability_score >= 60  # no two_column penalty
    # No two_column breakage
    types = [b.type for b in result.breakages]
    assert "two_column" not in types


def test_single_col_high_parsability_multiple_ats(ats_fixtures_dir: Path) -> None:
    """Single-column PDF → reasonably high parsability for several ATS."""
    from decroche.ats.parse_sim import parse_sim

    for ats_id in ("greenhouse", "icims", "taleo_oracle"):
        result = parse_sim(ats_fixtures_dir / "single_col.pdf", ats_id)
        assert result.parsability_score >= 50, f"Expected >=50 for {ats_id}, got {result.parsability_score}"
        types = [b.type for b in result.breakages]
        assert "two_column" not in types, f"Unexpected two_column breakage for {ats_id}"


def test_two_col_breakage_workday(ats_fixtures_dir: Path) -> None:
    """Two-column PDF → two_column breakage for workday (concatenate_lr)."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "two_col.pdf", "workday")
    types = [b.type for b in result.breakages]
    assert "two_column" in types, f"Expected two_column breakage, got {types}"
    # Parsability must be lower than single column
    # workday two_col fidelity 0.55 vs 0.95
    assert result.parsability_score < 80


def test_two_col_breakage_taleo(ats_fixtures_dir: Path) -> None:
    """Two-column PDF → two_column breakage for taleo_oracle."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "two_col.pdf", "taleo_oracle")
    types = [b.type for b in result.breakages]
    assert "two_column" in types


def test_two_col_breakage_icims(ats_fixtures_dir: Path) -> None:
    """Two-column PDF → two_column breakage for icims."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "two_col.pdf", "icims")
    types = [b.type for b in result.breakages]
    assert "two_column" in types


def test_table_breakage_lever(ats_fixtures_dir: Path) -> None:
    """DOCX with table → table breakage for lever (scramble)."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "table.docx", "lever")
    types = [b.type for b in result.breakages]
    assert "table" in types, f"Expected table breakage, got {types}"


def test_table_breakage_taleo(ats_fixtures_dir: Path) -> None:
    """DOCX with table → table breakage for taleo_oracle (worst scramble)."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "table.docx", "taleo_oracle")
    types = [b.type for b in result.breakages]
    assert "table" in types


def test_header_contact_workday(ats_fixtures_dir: Path) -> None:
    """DOCX with contact in header → header_contact breakage + contact in fields_lost for workday."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "header_contact.docx", "workday")
    types = [b.type for b in result.breakages]
    assert "header_contact" in types, f"Expected header_contact breakage, got {types}"
    assert "contact" in result.fields_lost, f"Expected contact in fields_lost, got {result.fields_lost}"


def test_unknown_ats_id_raises(ats_fixtures_dir: Path) -> None:
    """Unknown ATS id → ValueError with list of valid ids."""
    from decroche.ats.parse_sim import parse_sim

    with pytest.raises(ValueError, match="unknown_ats"):
        parse_sim(ats_fixtures_dir / "single_col.pdf", "unknown_ats")


def test_generic_fallback(ats_fixtures_dir: Path) -> None:
    """generic ATS id is accepted as valid."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "single_col.pdf", "generic")
    assert result.ats_id == "generic"
    assert 0 <= result.parsability_score <= 100


def test_ats_result_model_fields(ats_fixtures_dir: Path) -> None:
    """AtsParseResult has all required fields with correct types."""
    from decroche.ats.parse_sim import parse_sim
    from decroche.models import AtsParseResult

    result = parse_sim(ats_fixtures_dir / "single_col.pdf", "greenhouse")
    assert isinstance(result, AtsParseResult)
    assert isinstance(result.parsability_score, float)
    assert 0 <= result.parsability_score <= 100
    assert isinstance(result.fields_extracted, dict)
    assert isinstance(result.fields_lost, list)
    assert isinstance(result.breakages, list)
    assert result.fmt in ("pdf", "docx", "txt", "md", "unknown")


def test_parsability_score_clamped(ats_fixtures_dir: Path) -> None:
    """Parsability score is always in [0, 100]."""
    from decroche.ats.parse_sim import parse_sim

    for ats_id in ("workday", "taleo_oracle", "ashby", "generic"):
        result = parse_sim(ats_fixtures_dir / "two_col.pdf", ats_id)
        assert 0 <= result.parsability_score <= 100, f"Out of range for {ats_id}: {result.parsability_score}"


def test_breakage_has_required_fields(ats_fixtures_dir: Path) -> None:
    """Each Breakage has type, location, severity, fix."""
    from decroche.ats.parse_sim import parse_sim

    result = parse_sim(ats_fixtures_dir / "two_col.pdf", "workday")
    for b in result.breakages:
        assert b.type
        assert b.location
        assert b.severity in ("CRITICAL", "HIGH", "MEDIUM", "LOW")
        assert b.fix
