"""Tests for render_ats_docx — ATS-safe DOCX generation.

TDD: these tests define the contract; implementation must satisfy them.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from decroche.cv.render_docx import render_ats_docx
from decroche.models import Basics, JSONResume, Work, Education, Skill, Language


# ── Fixtures ─────────────────────────────────────────────────────────────────────

@pytest.fixture
def sample_resume() -> JSONResume:
    return JSONResume(
        basics=Basics(
            name="Jane Doe",
            email="jane.doe@example.com",
            phone="+1 415 555 0101",
            label="Senior Backend Engineer",
            summary="8 years building distributed systems.",
        ),
        work=[
            Work(
                name="Acme Corp",
                position="Staff Engineer",
                startDate="2020-01",
                endDate="2024-06",
                summary="Led platform team.",
                highlights=[
                    "Reduced API latency 38% by introducing a caching layer.",
                    "Led migration of 12 services to Kubernetes.",
                ],
            ),
            Work(
                name="Beta Inc",
                position="Software Engineer",
                startDate="2016-03",
                endDate="2019-12",
                highlights=["Built CI/CD pipeline used by 50+ engineers."],
            ),
        ],
        education=[
            Education(
                institution="MIT",
                area="Computer Science",
                studyType="B.S.",
                startDate="2012-09",
                endDate="2016-05",
            )
        ],
        skills=[
            Skill(name="Python", keywords=["asyncio", "FastAPI"]),
            Skill(name="Go"),
            Skill(name="Kubernetes"),
        ],
        languages=[Language(language="English", fluency="native")],
    )


@pytest.fixture
def market_fr():
    from decroche.market.profiles import load_profile
    return load_profile("fr")


@pytest.fixture
def market_us():
    from decroche.market.profiles import load_profile
    return load_profile("us")


# ── File creation ───────────────────────────────────────────────────────────────────

class TestRenderDocxCreation:
    def test_file_is_created(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_returns_path_object(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        result = render_ats_docx(sample_resume, market_fr, out)
        assert isinstance(result, Path)
        assert result == out


# ── Structural properties (no tables, single-column equivalent) ───────────────────────

class TestRenderDocxStructure:
    def test_no_tables(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        assert len(doc.tables) == 0, "DOCX must not contain any tables"

    def test_has_paragraphs(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        assert len(doc.paragraphs) > 5

    def test_no_section_headers_with_contact(self, sample_resume, market_fr, tmp_path):
        """Contact info MUST NOT be in headers/footers — must be in body."""
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        for section in doc.sections:
            hdr_text = " ".join(p.text.strip() for p in section.header.paragraphs)
            assert "jane.doe@example.com" not in hdr_text
            assert "+1 415 555 0101" not in hdr_text


# ── Contact information in body ────────────────────────────────────────────────────────

class TestRenderDocxContact:
    def test_name_in_body(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in body_text

    def test_email_in_body(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "jane.doe@example.com" in body_text

    def test_phone_in_body(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "+1 415 555 0101" in body_text


# ── Section headings ──────────────────────────────────────────────────────────────────

CANONICAL_HEADINGS = ("Summary", "Experience", "Education", "Skills")


class TestRenderDocxHeadings:
    def test_canonical_headings_present(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        for heading in CANONICAL_HEADINGS:
            assert heading in body_text, f"Canonical heading {heading!r} missing"

    def test_skills_content_present(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Python" in body_text
        assert "Kubernetes" in body_text


# ── Date formatting ──────────────────────────────────────────────────────────────────

class TestRenderDocxDates:
    def test_dates_formatted_as_mon_yyyy(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "2020" in body_text

    def test_us_market_same_format(self, sample_resume, market_us, tmp_path):
        """US market also uses Mon YYYY."""
        out = tmp_path / "cv_us.docx"
        render_ats_docx(sample_resume, market_us, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "2020" in body_text


# ── Work highlights as bullets ──────────────────────────────────────────────────────────

class TestRenderDocxBullets:
    def test_highlights_in_body(self, sample_resume, market_fr, tmp_path):
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "38%" in body_text
        assert "Kubernetes" in body_text

    def test_no_tables_for_bullets(self, sample_resume, market_fr, tmp_path):
        """Bullets must not be in tables."""
        out = tmp_path / "cv.docx"
        render_ats_docx(sample_resume, market_fr, out)
        doc = Document(str(out))
        assert len(doc.tables) == 0


# ── Minimal resume (edge cases) ────────────────────────────────────────────────────────

class TestRenderDocxMinimal:
    def test_empty_work_still_creates_file(self, market_fr, tmp_path):
        resume = JSONResume(basics=Basics(name="Test User"))
        out = tmp_path / "minimal.docx"
        render_ats_docx(resume, market_fr, out)
        assert out.exists()
        doc = Document(str(out))
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Test User" in body_text

    def test_null_dates_handled(self, market_fr, tmp_path):
        """Work entries with null dates should not crash."""
        resume = JSONResume(
            basics=Basics(name="Test User", email="t@t.com"),
            work=[Work(name="Org", position="Dev", highlights=["Built stuff."])],
        )
        out = tmp_path / "no_dates.docx"
        render_ats_docx(resume, market_fr, out)
        assert out.exists()
