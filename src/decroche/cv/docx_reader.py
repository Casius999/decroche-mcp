from __future__ import annotations

import io

from docx import Document


def read_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(c.text for c in row.cells))
    return "\n".join(lines)
