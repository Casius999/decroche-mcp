"""ATS-safe DOCX renderer.

Produces a single-column, table-free, header-free DOCX that round-trips
cleanly through parse_sim for workday and generic ATS profiles.

Rules (from ats_quirks.json + spec §5):
- Single column (DOCX is always treated as single-col by analyze_docx)
- No tables (they scramble in most ATS)
- Contact in body paragraphs (never in header/footer)
- Canonical section headings: Summary, Experience, Education, Skills,
  Certifications, Languages
- Date format: "Mon YYYY" (e.g. "Jan 2020")
- Bullets as plain paragraphs (no nested tables, no text boxes)
- Standard font: Calibri (ATS-safe)
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from decroche.models import JSONResume, MarketProfile

# Canonical month names (en)
_MONTHS = [
    "Jan", "Feb", "Mar", "Apr", "May", "Jun",
    "Jul", "Aug", "Sep", "Oct", "Nov", "Dec",
]

_DATE_RE = re.compile(r"^(\d{4})-(\d{1,2})(?:-\d+)?$")


def _fmt_date(raw: str | None) -> str:
    """Convert ISO date string to 'Mon YYYY' format.

    Examples:
        "2020-01" → "Jan 2020"
        "2020-01-15" → "Jan 2020"
        "2020" → "2020"
        None → ""
        "ongoing" → "Present"
    """
    if not raw:
        return ""
    raw = raw.strip()
    # Already looks like "Mon YYYY"
    if re.match(r"^[A-Za-z]{3}\s+\d{4}$", raw):
        return raw
    m = _DATE_RE.match(raw)
    if m:
        year = int(m.group(1))
        month = int(m.group(2))
        if 1 <= month <= 12:
            return f"{_MONTHS[month - 1]} {year}"
        return str(year)
    # e.g. "ongoing", "present", "current" → canonical "Present"
    if raw.lower() in ("ongoing", "present", "current", ""):
        return "Present"
    # Fallback: return as-is (year only, etc.)
    return raw


def _date_range(start: str | None, end: str | None) -> str:
    """Format a date range as 'Mon YYYY – Mon YYYY' or 'Mon YYYY – Present'."""
    s = _fmt_date(start)
    e = _fmt_date(end) or "Present"
    if s and e:
        return f"{s} – {e}"
    return s or e


def _set_font(run, size_pt: float = 11.0) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)


def _add_heading_paragraph(doc: Document, text: str, level: int = 1) -> None:
    """Add a section heading as a bold paragraph (no actual Heading style,
    which can confuse some ATS parsers). Uses bold + slightly larger font.

    Text is stored as-is (mixed-case canonical heading like "Summary") so ATS
    parsers recognise it. Visual weight comes from bold + spacing alone.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    _set_font(run, 11.0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _add_body_paragraph(doc: Document, text: str, indent: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, 10.5)
    if indent:
        p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)


def _add_bullet_paragraph(doc: Document, text: str) -> None:
    """Add a plain-text bullet (dash prefix, no table, no list style)."""
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    _set_font(run, 10.5)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)


def _add_divider(doc: Document) -> None:
    """Add a thin horizontal separator (underscores as plain text)."""
    p = doc.add_paragraph()
    run = p.add_run("─" * 60)
    _set_font(run, 6.0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def render_ats_docx(
    json_resume: JSONResume,
    market: MarketProfile,
    out_path: str | Path,
) -> Path:
    """Render an ATS-safe single-column DOCX from a JSONResume.

    Guarantees:
    - No tables
    - No header/footer content (contact in body)
    - Canonical section headings
    - Mon YYYY date format
    - Single-column layout
    - Standard Calibri font

    Args:
        json_resume: The structured resume data.
        market: Market profile (used for photo policy, etc.).
        out_path: Destination file path (will be created/overwritten).

    Returns:
        Path to the written DOCX file.
    """
    doc = Document()

    # ── Remove all header/footer content ────────────────────────────────────
    # Python-docx creates a default section; ensure header/footer are empty
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    # Clear header paragraphs
    for hdr_para in section.header.paragraphs:
        for run in hdr_para.runs:
            run.text = ""
        hdr_para.text = ""
    # Clear footer paragraphs
    for ftr_para in section.footer.paragraphs:
        for run in ftr_para.runs:
            run.text = ""
        ftr_para.text = ""

    # ── Set document-level margins (standard A4) ─────────────────────────────
    from docx.shared import Cm
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    basics = json_resume.basics

    # ── NAME (large, bold) ────────────────────────────────────────────────────
    name_text = basics.name or "Candidate"
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(name_text)
    run_name.bold = True
    _set_font(run_name, 16.0)
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_name.paragraph_format.space_after = Pt(2)

    # ── LABEL (job title) ─────────────────────────────────────────────────────
    if basics.label:
        p_label = doc.add_paragraph()
        run_label = p_label.add_run(basics.label)
        _set_font(run_label, 11.0)
        run_label.bold = False
        run_label.italic = True
        p_label.paragraph_format.space_after = Pt(2)

    # ── CONTACT INFO in body (NEVER in header) ────────────────────────────────
    contact_parts: list[str] = []
    if basics.email:
        contact_parts.append(basics.email)
    if basics.phone:
        contact_parts.append(basics.phone)
    if basics.url:
        contact_parts.append(basics.url)
    if basics.location and basics.location.city:
        loc_parts = [basics.location.city]
        if basics.location.region:
            loc_parts.append(basics.location.region)
        if basics.location.countryCode:
            loc_parts.append(basics.location.countryCode)
        contact_parts.append(", ".join(loc_parts))

    if contact_parts:
        p_contact = doc.add_paragraph()
        run_contact = p_contact.add_run(" | ".join(contact_parts))
        _set_font(run_contact, 10.0)
        p_contact.paragraph_format.space_after = Pt(4)

    # LinkedIn / social profiles
    for profile in basics.profiles:
        if profile.url:
            p_prof = doc.add_paragraph()
            label = profile.network or "Profile"
            run_prof = p_prof.add_run(f"{label}: {profile.url}")
            _set_font(run_prof, 10.0)
            p_prof.paragraph_format.space_after = Pt(1)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if basics.summary:
        _add_heading_paragraph(doc, "Summary")
        _add_divider(doc)
        _add_body_paragraph(doc, basics.summary)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    if json_resume.work:
        _add_heading_paragraph(doc, "Experience")
        _add_divider(doc)
        for job in json_resume.work:
            # Company | Position line
            company = job.name or ""
            position = job.position or ""
            date_range = _date_range(job.startDate, job.endDate)

            if company or position:
                line_parts: list[str] = []
                if position:
                    line_parts.append(position)
                if company:
                    line_parts.append(company)
                entry_line = " — ".join(line_parts) if line_parts else ""
                if date_range:
                    entry_line = f"{entry_line}  ({date_range})"
                p_job = doc.add_paragraph()
                run_job = p_job.add_run(entry_line)
                run_job.bold = True
                _set_font(run_job, 10.5)
                p_job.paragraph_format.space_before = Pt(4)
                p_job.paragraph_format.space_after = Pt(1)

            # Summary line
            if job.summary:
                _add_body_paragraph(doc, job.summary, indent=True)

            # Highlights as bullets
            for h in job.highlights:
                if h.strip():
                    _add_bullet_paragraph(doc, h.strip())

    # ── EDUCATION ────────────────────────────────────────────────────────────
    if json_resume.education:
        _add_heading_paragraph(doc, "Education")
        _add_divider(doc)
        for edu in json_resume.education:
            institution = edu.institution or ""
            area = edu.area or ""
            study_type = edu.studyType or ""
            date_range = _date_range(edu.startDate, edu.endDate)

            degree_parts: list[str] = []
            if study_type:
                degree_parts.append(study_type)
            if area:
                degree_parts.append(area)
            degree = ", ".join(degree_parts)

            line = institution
            if degree:
                line = f"{degree} — {institution}" if institution else degree
            if date_range:
                line = f"{line}  ({date_range})"

            if line.strip():
                p_edu = doc.add_paragraph()
                run_edu = p_edu.add_run(line.strip())
                _set_font(run_edu, 10.5)
                p_edu.paragraph_format.space_before = Pt(3)
                p_edu.paragraph_format.space_after = Pt(1)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    if json_resume.skills:
        _add_heading_paragraph(doc, "Skills")
        _add_divider(doc)
        # Flat list, comma-separated, NO table
        skill_names: list[str] = []
        for skill in json_resume.skills:
            if skill.name:
                skill_names.append(skill.name)
                if skill.keywords:
                    # Append keywords inline (e.g. "Python (asyncio, FastAPI)")
                    kw = ", ".join(k for k in skill.keywords if k)
                    skill_names[-1] = f"{skill.name} ({kw})"

        if skill_names:
            # Write as comma-separated body text (no table!)
            _add_body_paragraph(doc, ", ".join(skill_names))

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    # (not in JSONResume standard model, skip if absent)

    # ── LANGUAGES ─────────────────────────────────────────────────────────────
    if json_resume.languages:
        _add_heading_paragraph(doc, "Languages")
        _add_divider(doc)
        lang_parts: list[str] = []
        for lang in json_resume.languages:
            if lang.language:
                part = lang.language
                if lang.fluency:
                    part = f"{lang.language} ({lang.fluency})"
                lang_parts.append(part)
        if lang_parts:
            _add_body_paragraph(doc, ", ".join(lang_parts))

    # ── Write to disk ─────────────────────────────────────────────────────────
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
