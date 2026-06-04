"""Structural detectors for ATS parse simulation.

Operates on the ORIGINAL file (raw bytes / pdfplumber / python-docx),
returning a DocStructure dataclass that describes layout properties.
No LLM, no network, deterministic.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:(?:\+|00)\d{1,3}[\s.\-]?)?(?:\(?\d{1,4}\)?[\s.\-]?){2,5}\d{2,}")

# Words that strongly suggest contact info in a text block
_CONTACT_PATTERNS = re.compile(r"@|tel:|phone:|mobile:|fax:|\+\d|\d{2}[\s.\-]\d{2}", re.I)


@dataclass
class DocStructure:
    """Layout properties extracted from the raw document."""

    fmt: str  # "pdf" | "docx" | "txt" | "md"
    columns: int  # 1 or >=2
    has_tables: bool
    contact_in_header: bool
    page_count: int
    total_chars: int


# ── PDF detection ────────────────────────────────────────────────────────────────────────


def detect_columns(page) -> int:  # type: ignore[no-untyped-def]
    """Cluster word x0 positions → 1 vs ≥2 columns.

    Uses a simple gap-based approach: if the x-positions of words span
    more than 40% of the page width and there is a significant gap (>15%
    of page width) in the middle region, we declare two columns.
    """
    words = page.extract_words() or []
    if not words:
        return 1

    page_width = float(page.width) if page.width else 612.0
    x0_positions = [float(w["x0"]) for w in words]

    if not x0_positions:
        return 1

    x_min = min(x0_positions)
    x_max = max(x0_positions)
    span = x_max - x_min

    if span < page_width * 0.3:
        # All text in a narrow band → single column
        return 1

    # Look for a gap in the middle 30–70% region of the page
    mid_lo = page_width * 0.30
    mid_hi = page_width * 0.70

    # Histogram: does any horizontal band in the middle have NO words?
    # We use a coarser approach: sort x0s and find the largest gap in middle region
    mid_xs = sorted(x for x in x0_positions if mid_lo <= x <= mid_hi)
    if not mid_xs:
        # No words in the middle → strong sign of two columns with a gutter
        return 2

    # Check if there are clusters on both sides of the midpoint with a gap
    left_count = sum(1 for x in x0_positions if x < page_width * 0.40)
    right_count = sum(1 for x in x0_positions if x > page_width * 0.50)

    if left_count >= 2 and right_count >= 2:
        # Significant content on both sides → two columns
        # Confirm by checking gap between the two clusters
        left_max = max((x for x in x0_positions if x < page_width * 0.40), default=0)
        right_min = min((x for x in x0_positions if x > page_width * 0.50), default=page_width)
        gap = right_min - left_max
        if gap > page_width * 0.10:
            return 2

    return 1


def has_tables_pdf(page) -> bool:  # type: ignore[no-untyped-def]
    """Return True if pdfplumber finds table-like structures on the page."""
    tables = page.find_tables()
    return len(tables) > 0


def _contact_text(text: str) -> bool:
    """Return True if text contains email or phone contact info."""
    return bool(EMAIL_RE.search(text) or _CONTACT_PATTERNS.search(text))


def header_footer_contact_pdf(pdf) -> bool:  # type: ignore[no-untyped-def]
    """Check if email/phone appears in the top ~12% or bottom ~12% Y-band."""
    for page in pdf.pages:
        h = float(page.height) if page.height else 792.0
        top_band = h * 0.12
        bottom_band = h * 0.88

        words = page.extract_words() or []
        for w in words:
            y0 = float(w.get("top", h / 2))
            if y0 <= top_band or y0 >= bottom_band:
                if _contact_text(w.get("text", "")):
                    return True

        # Also check page-level text in those bands
        cropped_top = page.crop((0, 0, page.width, top_band))
        top_text = cropped_top.extract_text() or ""
        if _contact_text(top_text):
            return True

        cropped_bottom = page.crop((0, bottom_band, page.width, page.height))
        bottom_text = cropped_bottom.extract_text() or ""
        if _contact_text(bottom_text):
            return True

    return False


def analyze_pdf(data: bytes) -> DocStructure:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        total_chars = 0
        max_cols = 1
        any_table = False

        for page in pdf.pages:
            text = page.extract_text() or ""
            total_chars += len(text)
            cols = detect_columns(page)
            if cols > max_cols:
                max_cols = cols
            if has_tables_pdf(page):
                any_table = True

        contact_hf = header_footer_contact_pdf(pdf)

    return DocStructure(
        fmt="pdf",
        columns=max_cols,
        has_tables=any_table,
        contact_in_header=contact_hf,
        page_count=page_count,
        total_chars=total_chars,
    )


# ── DOCX detection ──────────────────────────────────────────────────────────────────────


def _header_footer_text_docx(doc) -> str:  # type: ignore[no-untyped-def]
    """Extract text from all section headers and footers."""
    parts: list[str] = []
    for section in doc.sections:
        try:
            hdr = section.header
            if hdr:
                parts.append(" ".join(p.text for p in hdr.paragraphs))
        except Exception:  # noqa: BLE001
            pass
        try:
            ftr = section.footer
            if ftr:
                parts.append(" ".join(p.text for p in ftr.paragraphs))
        except Exception:  # noqa: BLE001
            pass
    return " ".join(parts)


def analyze_docx(data: bytes) -> DocStructure:
    import io
    from docx import Document

    doc = Document(io.BytesIO(data))

    total_chars = sum(len(p.text) for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_chars += len(cell.text)

    has_tbl = len(doc.tables) > 0

    hf_text = _header_footer_text_docx(doc)
    contact_hf = _contact_text(hf_text)

    return DocStructure(
        fmt="docx",
        columns=1,  # DOCX is always treated as single-column (tables detected separately)
        has_tables=has_tbl,
        contact_in_header=contact_hf,
        page_count=1,  # python-docx does not expose page count easily
        total_chars=total_chars,
    )


# ── Plain text / markdown ────────────────────────────────────────────────────────────────────


def analyze_text(data: bytes, fmt: str = "txt") -> DocStructure:
    text = data.decode("utf-8", errors="replace")
    return DocStructure(
        fmt=fmt,
        columns=1,
        has_tables=False,
        contact_in_header=False,
        page_count=1,
        total_chars=len(text),
    )


# ── Dispatcher ─────────────────────────────────────────────────────────────────────────────


def analyze_file(path: str | Path, data: bytes | None = None) -> DocStructure:
    """Analyze a document and return its DocStructure."""
    p = Path(path)
    ext = p.suffix.lower()
    raw = data if data is not None else p.read_bytes()

    if ext == ".pdf":
        return analyze_pdf(raw)
    if ext == ".docx":
        return analyze_docx(raw)
    if ext in (".txt", ".md"):
        return analyze_text(raw, fmt=ext.lstrip("."))
    # Unknown format — minimal analysis
    return DocStructure(
        fmt="unknown",
        columns=1,
        has_tables=False,
        contact_in_header=False,
        page_count=1,
        total_chars=len(raw),
    )
